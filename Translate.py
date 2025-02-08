from docx import Document
from deep_translator import GoogleTranslator
from fpdf import FPDF
import os

def translate_text(text, source_lang='en', target_lang='es'):
    translator = GoogleTranslator(source=source_lang, target=target_lang)
    return translator.translate(text)

def translate_docx(input_path, output_path, output_format='docx'):
    # Cargar documento Word
    doc = Document(input_path)
    
    # Traducir cada párrafo
    translated_doc = Document()
    for para in doc.paragraphs:
        translated_text = translate_text(para.text)
        translated_doc.add_paragraph(translated_text)
    
    # Guardar en el formato deseado
    if output_format == 'docx':
        translated_doc.save(output_path)
    elif output_format == 'pdf':
        pdf = FPDF()
        pdf.set_auto_page_break(auto=True, margin=15)
        pdf.add_page()
        pdf.set_font("Arial", size=12)
        
        for para in translated_doc.paragraphs:
            pdf.multi_cell(0, 10, para.text)
            pdf.ln()
        
        pdf.output(output_path)
    else:
        raise ValueError("Formato de salida no soportado. Use 'docx' o 'pdf'.")

if __name__ == "__main__":
    input_file = "C:\Users\jorge\Downloads\kl_047.12.6_kaspersky_next_edr_optimum_en_unit1_v1.0.docx"  # Especifica aquí la ruta del archivo a traducir
    output_file = "C:\Users\jorge\Downloads\out.docx"  # Especifica aquí la ruta del archivo de salida
    output_format = output_file.split('.')[-1]
    
    translate_docx(input_file, output_file, output_format)
    print(f"Traducción completada. Archivo guardado en: {output_file}")